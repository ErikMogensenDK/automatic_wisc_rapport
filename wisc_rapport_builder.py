import pandas as pd
from texts.anbefalinger import *
from texts.general_intro import*
from docx import Document
from docx.shared import Inches
from datetime import datetime

class Builder:
	#def __init__(self, input_dict):
	#	#self.result_df = input_dataframe
	#	self.input_dict = input_dict
	def __init__(self, input_path):
		self.result_df = pd.read_excel(input_path)


#	def create_score_df(scores):
#		confidence_intervals = calculate_confidence_intervals(scores)
#		
#		return score_df
#	def calculate_confidence_intervals(scores):
#		for score in scores:
#
#		return confidence_intervals
#	
#	def calculate_confidence_interval:
#		ci_upper = (100+3.92)
#		return ci
	def create_df_from_dict(result):
		df_from_dict = pd.DataFrame.from_dict
		return df_from_dict

	def get_introduction(self):
		introduction = generel_intro
		return introduction
	
	def get_comparison_to_mean(self, score):
		if score < 70:
			return 'Langt under gennemsnittet'
		if score > 69 and score < 85:
			return 'Noget under gennemsnittet'
		if score > 84 and score < 90:
			return 'Gennemsnittets nederste del'
		if score > 90 and score < 109:
			return 'Gennemsnitligt'
		if score > 108 and score <115:
			return 'Øverste del af gennemsnittet'
		if score > 114 and score <130:
			return 'Noget over gennemsnittet'
		if score > 129: 
			return 'Langt over gennemsnittet'

	def create_description_of_scores(self, result_dict):
		long_version_dict = {'VFI': 'Verbalt Forståelses-Indeks', 
		       'HIK': 'Hele skalaen IntelligensKvotient', 
		       'VSI': 'VisuoSpatialt Indeks', 
		       'FHI': 'ForarbejdningsHastigheds-Indeks', 
		       'AHI': 'ArbejdsHukommelses-Indeks',
		       'RSI': 'Logisk Ræsonnerings-Indeks'}

		description = ''
		for result in result_dict:
			comparison_to_mean = self.get_comparison_to_mean(result['score'])
			description = description + f'''{result['mål']} ({long_version_dict[result['mål']]}) blev målt til {result['score']} (95% KI mellem {result['95%ki']}), hvilket er {comparison_to_mean}. Denne score var {result['percentil']}. percentil, hvilket vil sige at {result['percentil']}% af børnene i norm-gruppen scorede lavere. '''
			#descriptions.append(description)
		self.long_version_dict = long_version_dict
		return description
	
	def add_table_to_document(self, document, result_dict):
		table = document.add_table(rows=7, cols=5) 
		hdr_cells = table.rows[0].cells 
		# indsæt evt. også lang version?
		#hdr_cells[0].text = 'Indeks' + self.long_version_dict()
		hdr_cells[0].text = 'Indeks' 
		hdr_cells[1].text = 'Score' 
		hdr_cells[2].text = '95%KI' 
		hdr_cells[3].text = 'Percentil' 
		hdr_cells[4].text = 'Beskrivelse' 
		for item in result_dict: 
			row_cells = table.add_row().cells 
			row_cells[0].text = str(item['mål'])
			row_cells[1].text = str(item['score']) 
			row_cells[2].text = str(item['95%ki'])
			row_cells[3].text = str(item['percentil'])
			row_cells[4].text = self.get_comparison_to_mean(item['score'])
		return document

	
	def get_recommendations_lav(self, result_dict):
		recommendations = []
		for result in result_dict:
			if result['score'] < 86:
				recommendation = self.get_recommendation_of_specific_index(result['mål'])
				recommendations.append(recommendation)
		recommendations = '\n'.join(recommendations)
		return recommendations
	
	def get_recommendation_of_specific_index(self, index):
		match index:
			case 'VFI':
				return VFI_anbefaling_lav
			case 'HIK':
				return HIK_anbefaling_lav 
			case 'AHI':
				return AHI_anbefaling_lav 
			case 'FHI':
				return FHI_anbefaling_lav 
			case 'VSI':
				return VFI_anbefaling_lav
			case 'RSI':
				return RSI_anbefaling_lav

	def get_result(self, result_df):
		indexes = ['VFI', 'VSI', 'RSI', 'AHI', 'FHI', 'HIK']
		filtered_df = result_df[result_df.iloc[:, 0].isin(indexes)]
		result = filtered_df.to_dict(orient='records')
		return result

	def get_recommendations_for_HIK(self, result):
		# tjek om HIK er under eller over 85 og indsæt evt. følgende
		if result[0]['score'] < 85:
			return fortsat_intro_til_lav_begavelse
		if result[0]['score'] > 115:
			return fortsat_intro_til_høj_begavelse



	def build_rapport(self, save_name):
		document = Document()
		document.add_heading(f'WISC-IV rapport {datetime.today().strftime("%d-%m-%y")}', 0)
		document.add_paragraph(generel_intro)
		document.add_heading('Testresultat', 1)
		result = self.get_result(self.result_df)
		# testing if i can just pass dict:
		#result = self.input_dict
		document.add_paragraph(self.create_description_of_scores(result))

		document = self.add_table_to_document(document, result)

		document.add_heading('Klinisk Indtryk', 1)
		document.add_paragraph()


		document.add_heading('Anbefalinger', 1)
		document.add_paragraph(self.get_recommendations_for_HIK(result))
		document.add_paragraph(self.get_recommendations_lav(result))
		document.add_paragraph(generelle_anbefalinger)
		document.save(f'{save_name}.docx')
		return document


	#def create_description()
		
# Insert general description

# create df from excel
# for row 2-6 (or whatever) in excel: report each col from df


# Insert general recommendations
# create recommendations based on scores
# for each column containing scores, if score lower than 85: 
#	Insert recommendation based on category in row with score lower than 85

# Create summary?