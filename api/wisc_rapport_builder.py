import pandas as pd
from texts.anbefalinger import *
from texts.general_intro import*
from docx import Document
from docx.shared import Inches
from datetime import datetime
import numpy as np

class Builder:
	def __init__(self, input_dict):
		self.input_dict = input_dict

	def get_comparison_to_mean(self, score):
		if score < 70:
			return 'Langt under gennemsnittet'
		if score > 69 and score < 85:
			return 'Noget under gennemsnittet'
		if score > 84 and score < 93:
			return 'Gennemsnittets nederste del'
		if score > 92 and score < 108:
			return 'Gennemsnitligt'
		if score > 107 and score <116:
			return 'Gennemsnittets øverste del'
		if score > 115 and score <130:
			return 'Noget over gennemsnittet'
		if score > 129: 
			return 'Langt over gennemsnittet'

	def create_description_of_scores(self, result_dict):
		long_version_dict = {'VFI': 'verbalt forståelses-indeks', 
		       'HIK': 'hele skalaen intelligenskvotient', 
		       'VSI': 'visuospatialt indeks', 
		       'FHI': 'forarbejdningshastigheds-indeks', 
		       'AHI': 'arbejdshukommelses-indeks',
		       'RSI': 'logisk ræsonnerings-indeks'}

		description = ''
		for result in result_dict:
			comparison_to_mean = self.get_comparison_to_mean(result['score'])
			description = description + f'''{result['mål']} ({long_version_dict[result['mål']]}) blev målt til {int(result['score'])} (95% KI mellem {int(result['95%kil'])}-{int(result['95%kiu'])}), hvilket er {comparison_to_mean.lower()}. Denne score var {result['percentil']}. percentil, hvilket vil sige at {result['percentil']}% af børnene i norm-gruppen scorede lavere. '''
		self.long_version_dict = long_version_dict
		return description

	def add_table_to_document(self, document, result_dict):
		table = document.add_table(rows=1, cols=5) 
		table.style = "Light List"
		hdr_cells = table.rows[0].cells 
		hdr_cells[0].text = 'Indeks' 
		hdr_cells[1].text = 'Score' 
		hdr_cells[2].text = 'Percentil' 
		hdr_cells[3].text = '95%KI' 
		hdr_cells[4].text = 'Beskrivelse' 
		for item in result_dict: 
			row_cells = table.add_row().cells 
			row_cells[0].text = str(item['mål'])
			row_cells[1].text = str(int(item['score'])) 
			row_cells[2].text = str(item['percentil'])
			text_for_ki = str(int(item['95%kil'])) + '-' + str(int(item['95%kiu']))
			row_cells[3].text = text_for_ki
			row_cells[4].text = self.get_comparison_to_mean(item['score'])
		return document

	
	def get_recommendations_lav(self, result_dict):
		recommendations = []
		for result in result_dict:
			if result['score'] < 86:
				recommendation = self.get_recommendation_of_specific_index(result['mål'])
				recommendations.append(recommendation)
		return recommendations

	def get_recommendations_lav_headings(self, result_dict):
		recommendations_headers = []
		for result in result_dict:
			if result['score'] < 86:
				header = self.get_header_of_specific_index(result['mål'])
				recommendations_headers.append(header)
		return recommendations_headers
	
	def get_recommendation_of_specific_index(self, index):
		match index:
			case 'VFI':
				return VFI_anbefaling_lav
			case 'HIK':
				return HIK_anbefaling_lav
			case 'AHI':
				return AHI_anbefaling_lav 
			case 'FHI':
				return FHI_anbefaling_lav 
			case 'VSI':
				return VFI_anbefaling_lav
			case 'RSI':
				return RSI_anbefaling_lav

	def get_header_of_specific_index(self, index):
		match index:
			case 'VFI':
				return "Anbefalinger til støtte af nedsat verbal forståelse"
			case 'HIK':
				return "\n"
			case 'AHI':
				return "Anbefalinger til støtte af nedsat arbejdshukommelse"
			case 'FHI':
				return "Anbefalinger til støtte af nedsat forarbejdningshastighed"
			case 'VSI':
				return "Anbefalinger til støtte af nedsat visuospatial bearbejdning"
			case 'RSI':
				return "Anbefalinger til støtte af nedsat ræsonnering"

	def get_recommendations_for_HIK(self, result):
		# Check if IQ over or under average and insert recommendations if that should be the case.
		if result[0]['score'] < 85:
			return fortsat_intro_til_lav_begavelse
		if result[0]['score'] > 115:
			return fortsat_intro_til_høj_begavelse

	def build_rapport(self, save_name):
		document = Document()
		document.add_heading(f'WISC-V rapport {datetime.today().strftime("%d-%m-%y")}', 0)
		document.add_heading(f'Generel beskrivelse af den kognitive test WISC-V', 1)
		document.add_paragraph(generel_intro)
		document.add_heading(f'Beskrivelse af forskellige indekser', 1)
		document.add_paragraph(beskrivelse_af_indekser)
		document.add_heading('Testresultat', 1)

		result = self.input_dict
		document.add_paragraph(self.create_description_of_scores(result))

		document = self.add_table_to_document(document, result)

		document.add_heading('Klinisk Indtryk', 1)
		document.add_paragraph()


		document.add_heading('Anbefalinger', 1)
		document.add_paragraph(self.get_recommendations_for_HIK(result))

		document.add_heading('Generelle anbefalinger')
		document.add_paragraph(generelle_anbefalinger)

		recommendations = self.get_recommendations_lav(result)
		recommendations_headers = self.get_recommendations_lav_headings(result)
		for i in range(len(recommendations)):
			document.add_heading(recommendations_headers[i])
			document.add_paragraph(recommendations[i])
		document.save(f'{save_name}.docx')
		return document

